#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sun Aug 22 21:53:57 2021

@author: ricardobortothopker
"""

import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from math import pi,inf,exp,sqrt
from random import shuffle
import operator
from complexFuncs import complexity,DSM_rearrange
import pickle
from scipy.spatial import ConvexHull, convex_hull_plot_2d
from TSP_Functions import calculate_euclidean_distance_matrix,SA_TSP,total_distance,NeighborFunc,plotResults,cal_tsp_complexity, create_TSP_hull, create_TSP_pareto
rng = np.random.default_rng()
from TSP_Functions import *
# def run_t():
url = r'/Users/ricardobortothopker/OneDrive - Massachusetts Institute of Technology/Classes/Thesis/excels/Points for TSP/'
name=url+'TSP Problems3_corrected.pkl'
def create_problems_tsp(name=url+r'point count for TSP3.csv'):
    # url2 = r'point count for TSP.csv'
    df = pd.read_csv(name)
    outside = df['outside points'].to_list()
    total = df['points'].to_list()
    totOut = list(zip(total,outside))
    xyArrTot =[]
    for tot,out in totOut:
        
        xyArrTot.append(create_TSP_hull(out,tot))
        print(xyArrTot[-1])
    return xyArrTot
def xyArrTot_toDict(xyArrTot,store_dict={},count = 0):
    for i in xyArrTot:
        store_dict['id '+str(count)]=i
        count += 1
    return store_dict
def save_TSP_Problems(xyArrTot,name=url+'TSP Problems3.pkl',store_dict={},count = 0):
    store_dict = xyArrTot_toDict(xyArrTot,store_dict,count)
    # for i in xyArrTot:
    #     store_dict['id '+str(count)]=i
    #     count += 1
    with open(name, 'wb') as f:  # Python 3: open(..., 'wb')
        pickle.dump(store_dict, f)
def load_TSP_Problems(url=url+'TSP Problems3.pkl'):
    with open(url,'rb') as f:  # Python 3: open(..., 'rb')
        xyArrDict = pickle.load(f)
    return xyArrDict
def update_TSP_problems(xyArrDict, newxyArrTot,url = url+'TSP Problems3.pkl'):
    keys = list(xyArrDict.keys())
    ids = []
    for i in range(len(keys)):
        ids.append(int(keys[i][3:]))
    lastKey = sorted(ids)[-1]+1
    save_TSP_Problems(newxyArrTot,url,xyArrDict,lastKey)
from operator import sub
def get_aspect(ax):
    # Total figure size
    figW, figH = ax.get_figure().get_size_inches()
    # Axis size on figure
    _, _, w, h = ax.get_position().bounds
    # Ratio of display units
    disp_ratio = (figH * h) / (figW * w)
    # Ratio of data units
    # Negative over negative because of the order of subtraction
    data_ratio = sub(*ax.get_ylim()) / sub(*ax.get_xlim())

    return disp_ratio / data_ratio
def visualize_TSP(xyArrDict,show_hull=True,show=True,corrected = True,equal=True):
    count =0
    newDict = {}
    for key, i in xyArrDict.items():
        if key =='id 29':
            continue
        x = i[:,0]
        y = i[:,1]
        fig,ax = plt.subplots()
        ax.scatter(x,y)
        if show_hull:
            hull = ConvexHull(i)
            hullVert = i[np.append(hull.vertices,hull.vertices[0])]
            ax.plot(hullVert[:,0],hullVert[:,1],'--',c='r')
        ax.set_xticks([])
        ax.set_yticks([])
        ax.annotate('Test {}'.format(key),xy=(0.00, 1.02), xycoords='axes fraction')
        if equal:
            ax.set_aspect('equal')
        print(ax.get_aspect())
        print(get_aspect(ax))
        if show:
            fig.show()
        else:
            fig.savefig('TSP {}.png'.format(count))
            fig.show()
        count+=1
        arr = i.copy()
        arr[:,0] *= 1/get_aspect(ax)
        newDict[key] = arr
    if corrected:
        return newDict
probs = load_TSP_Problems()    
new_probs = visualize_TSP(probs,show_hull=False,show=True,corrected =True,equal=False)
def run_solve_TSP(xyArrDict,visualize_unsolved_TSP = True,visualize_solved_TSP = True):
    best =[]
    accept=[]
    sols=[]
    it =[]
    bestSol=[]
    totComplexity=[]
    visualize_TSP(xyArrDict,visualize_unsolved_TSP)
    for key, i in xyArrDict.items():
        
        dist_matrix = calculate_euclidean_distance_matrix(i)
        init = np.array(range(len(i)))
        shuffle(init)
        sol,best_hist,it_hist,accep_hist = SA_TSP(total_distance, NeighborFunc, init, [dist_matrix], [.2],[3,0,.95,200,200,2])
        if visualize_solved_TSP:
            plotResults(sol,i,dist_matrix)
        where0 = np.where(sol==0)[0][0]
        sol = np.append(sol[where0:],sol[:where0])
        sols.append(sol)
        bestSol.append(total_distance(sol,dist_matrix))
        best.append(best_hist)
        accept.append(accep_hist)
        it.append(it_hist)
        totComplexity.append(cal_tsp_complexity(i))
    return sols,bestSol,best,accept,it,totComplexity
def create_b_dict(xyArrDict,sols,bestSol,best,accept,it,totComplexity):
    out_dict = {}
    i=0
    for name in xyArrDict.keys():
        # name = 'id ' +str(i)
        out_dict[name] ={}
        out_dict[name]['best']=best[i]
        out_dict[name]['bestSol']=bestSol[i]
        # out_dict[name]['it']=it[i]
        # out_dict[name]['accept']=accept[i]
        out_dict[name]['sol']=sols[i]
        out_dict[name]['xyArr']=xyArrDict[name]
        out_dict[name]['complexity']=totComplexity[i]
        i+=1
    return out_dict
def save_BSF(xyArrDict,sols,bestSol,best,accept,it,totComplexity,url=url+'TSP BSF3.pkl'):
    out_dict = create_b_dict(xyArrDict,sols,bestSol,best,accept,it,totComplexity)
    with open(url, 'wb') as f:  # Python 3: open(..., 'wb')
        pickle.dump(out_dict, f)
def load_BSF(name=url+r'TSP BSF3.pkl'):
    with open(name,'rb') as f:  # Python 3: open(..., 'rb')
        sol_dict = pickle.load(f)
    return sol_dict      
def update_BSF(newDict,oldDict,url = url+'TSP BSF3.pkl'):
    out_dict =oldDict
    oKeys = oldDict.keys()
    for nKey in newDict.keys():
        if nKey in oKeys:
            if newDict[nKey]['bestSol']<oldDict[nKey]['bestSol']:
                out_dict[nKey] =newDict[nKey]
                print(nKey+' improved')
        else:
            out_dict[nKey] =newDict[nKey]
    with open(url, 'wb') as f:  # Python 3: open(..., 'wb')
        pickle.dump(out_dict, f)
    return out_dict
def plot_bsf(bsfDict):
    for key in bsfDict.keys():
        
        sol = bsfDict[key]['sol']
        xy = bsfDict[key]['xyArr']
        where0 = np.where(sol==0)[0][0]
        sol = np.append(sol[where0:],sol[:where0])
        plotResults(sol,xy,None,key+'\nSol: '+ str(sol)+'\n',True)
def summary_to_csv(bsf,url=url+r'TSP summary.csv'):
    df = pd.DataFrame(bsf).T
    xyArrTot = df['xyArr'].to_list()
    Cdict = cal_tsp_complexityMult(xyArrTot)
    df2 = pd.DataFrame(Cdict,index=df.index)
    df3 = df[['bestSol','sol']]
    out = pd.concat([df3,df2[['C','C1','C2','C3']]],axis=1)
    out.to_csv(url)
def create_and_visualize_TSP():
    xyArrTot = create_problems_tsp()
    xyArrDict = xyArrTot_toDict(xyArrTot)
    
    visualize_TSP(xyArrDict)
    return xyArrTot
def select_problems_to_run(lst,xyArrDict):
    outDict = {}
    for i in lst:
        outDict['id '+str(i)]= xyArrDict['id '+str(i)]
    return outDict
def create_experiments():
    from itertools import combinations
    data = list(range(0,30,2))
    cc = np.array(list(combinations(data,13)))
    np.random.seed(22689)
    random = np.random.randint(2,size=np.shape(cc))
    notRandom = random==0
    problems = cc+random
    problems2 = cc+notRandom
    problemsTot = np.concatenate([problems.T,problems2.T],axis=1).T
    return problemsTot
def write_problem_to_word(path = r'/Users/ricardobortothopker/OneDrive - Massachusetts Institute of Technology/Classes/Thesis/TSP to solve/TSP Problems/'):
    import docx
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import random
    problemsTot = create_experiments()
    mydoc = docx.Document()
    countTot = 0
    for problem in problemsTot:
        countTot +=1
        count = 0
        probCopy = problem.copy()
        random.seed(countTot)
        shuffle(probCopy)
        for prob in probCopy:
            count +=1
            imageP = path+'TSP {}.png'.format(prob)
            mydoc.add_paragraph('                              Problem  {} from test {}'.format(count,countTot))
            mydoc.add_picture(imageP, width=docx.shared.Inches(7.5), height=docx.shared.Inches(5))
            mydoc.add_page_break()
    sections = mydoc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "Canonical Experiment in System Complexity and Human Effort"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mydoc.save(path+r'TSP Problems.docx')
#newArr= create_problems_tsp()
# xyArrDict = load_TSP_Problems()
# visualize_TSP(xyArrDict,False,False)
#write_problem_to_word()
# # # xyArrDict = select_problems_to_run(list(range(16,32)),xyArrDict)
# xyArrDict = select_problems_to_run(list(range(32,48)),xyArrDict)
# # # # visualize_TSP(xyArrDict,True,False)
# # # xyArrDict= {'id 14':testDict['id 14']}
# sols,bestSol,best,accept,it,totComplexity = run_solve_TSP(xyArrDict,False,False)
# bsf = load_BSF()
# newDict = create_b_dict(xyArrDict,sols,bestSol,best,accept,it,totComplexity)
# bsf = update_BSF(newDict,bsf)
# plot_bsf(bsf)
# xyArrTotNew = create_and_visualize_TSP()
# update_TSP_problems(xyArrDict, xyArrTotNew)
# summary_to_csv(load_BSF())
